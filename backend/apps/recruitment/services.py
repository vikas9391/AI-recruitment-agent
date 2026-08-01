from __future__ import annotations

import json
import logging
import os
import re
from typing import Any, Dict, List

import fitz  # PyMuPDF
from django.conf import settings
from docx import Document as DocxDocument
from django.db import transaction
from django.db.models import Q
from rest_framework.exceptions import NotFound, ValidationError


from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from .models import (
        Application,
        ApplicationHistory,
        Candidate,
        CandidateDocument,
        Resume,
        ResumeAnalysis,
        ResumeScore,
    )

logger = logging.getLogger(__name__)


class ResumeParserService:
    """Extracts raw text from uploaded resume files (PDF / DOCX)."""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        text_parts: List[str] = []
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    text_parts.append(page.get_text())
        except Exception as exc:
            logger.exception("Failed to parse PDF resume: %s", file_path)
            raise ValidationError(f"Unable to parse PDF resume: {exc}")
        return "\n".join(text_parts).strip()

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            return "\n".join(paragraphs).strip()
        except Exception as exc:
            logger.exception("Failed to parse DOCX resume: %s", file_path)
            raise ValidationError(f"Unable to parse DOCX resume: {exc}")

    @staticmethod
    def extract_text(file_path: str, file_type: str) -> str:
        file_type = file_type.lower()
        if file_type == "pdf":
            text = ResumeParserService.extract_text_from_pdf(file_path)
        elif file_type == "docx":
            text = ResumeParserService.extract_text_from_docx(file_path)
        else:
            raise ValidationError(f"Unsupported resume file type: {file_type}")

        if not text or len(text.strip()) < 20:
            raise ValidationError("Resume appears to be empty or unreadable.")
        return text


class GeminiResumeAnalyzer:
    """Sends resume text + job context to Google Gemini and returns structured analysis.

    Gemini is responsible for generating ALL scoring sub-components
    (skills, experience, education, projects, certifications, ats, keyword match).
    The backend only validates the structure/ranges and computes the final
    weighted overall score — no scoring heuristics are hardcoded server-side.
    """

    REQUIRED_KEYS = [
        "skills", "matched_skills", "missing_skills",
        "experience", "education", "projects", "certifications",
        "skills_match_score", "experience_score", "education_score",
        "projects_score", "certifications_score", "keyword_match", "ats_score",
        "strengths", "weaknesses", "improvement_report",
    ]

    SCORE_FIELDS = [
        "skills_match_score", "experience_score", "education_score",
        "projects_score", "certifications_score", "keyword_match", "ats_score",
    ]

    LIST_FIELDS = [
        "skills", "matched_skills", "missing_skills", "projects",
        "certifications", "strengths", "weaknesses", "improvement_report",
    ]

    # Generous cap purely to protect against pathological/corrupted files,
    # not a realistic resume-length restriction. Gemini 1.5 models support
    # very large context windows, so this is intentionally high.
    MAX_RESUME_CHARS = 60000

    def __init__(self) -> None:
        api_key = getattr(settings, "GEMINI_API_KEY", None) or os.environ.get("GEMINI_API_KEY")
        if not api_key:
            raise ValidationError("GEMINI_API_KEY is not configured on the server.")

        import google.generativeai as genai

        genai.configure(api_key=api_key)
        model_name = getattr(settings, "GEMINI_MODEL_NAME", "gemini-1.5-flash")
        self.model = genai.GenerativeModel(model_name)

    def _build_prompt(self, resume_text: str, job: Any) -> str:
        skills_required = job.skills_required if isinstance(job.skills_required, list) else []
        truncated_resume = resume_text[: self.MAX_RESUME_CHARS]

        return f"""
You are an expert technical recruiter and resume screening AI.

Analyze the RESUME TEXT against the JOB DETAILS below and respond with ONLY valid JSON.
Do not include markdown code fences, explanations, or any text outside the JSON object.

JOB DETAILS:
Title: {job.title}
Department: {job.department}
Employment Type: {job.employment_type}
Experience Required: {job.experience_required}
Education Required: {job.education_required}
Required Skills: {", ".join(skills_required) if skills_required else "Not specified"}
Job Description: {job.description}
Requirements: {job.requirements}
Responsibilities: {job.responsibilities}

RESUME TEXT:
\"\"\"
{truncated_resume}
\"\"\"

Return a JSON object with EXACTLY this structure. All *_score fields and
keyword_match/ats_score MUST be integers from 0 to 100, based on YOUR expert
judgement of how well the resume satisfies that specific dimension for THIS job:

{{
  "skills": ["list of all skills found in the resume"],
  "matched_skills": ["skills from resume that match the job's required skills"],
  "missing_skills": ["required job skills NOT found in the resume"],
  "experience": "short summary of candidate's relevant work experience",
  "education": "short summary of candidate's education background",
  "projects": ["list of notable projects mentioned"],
  "certifications": ["list of certifications mentioned"],

  "skills_match_score": <integer 0-100, how well candidate's skills cover the job's required skills>,
  "experience_score": <integer 0-100, relevance and sufficiency of candidate's experience for this role>,
  "education_score": <integer 0-100, how well candidate's education matches the job's education requirement>,
  "projects_score": <integer 0-100, quality/relevance of candidate's projects to this role>,
  "certifications_score": <integer 0-100, relevance/value of candidate's certifications to this role>,
  "keyword_match": <integer 0-100, how well resume keywords match the job description>,
  "ats_score": <integer 0-100, how well-formatted/ATS-friendly the resume is>,

  "strengths": ["list of candidate strengths relevant to this job"],
  "weaknesses": ["list of candidate weaknesses/gaps relevant to this job"],
  "improvement_report": ["specific actionable suggestions to improve this resume for this job"]
}}

Scoring guidance:
- Score strictly relative to THIS job's requirements, not in the abstract.
- If information for a dimension (e.g. certifications) is entirely absent, score it low rather than omitting it.
- Be consistent and evidence-based; do not inflate scores.

Respond with ONLY the JSON object.
""".strip()

    def _parse_response(self, raw_text: str) -> Dict[str, Any]:
        cleaned = raw_text.strip()
        cleaned = re.sub(r"^```json\s*|^```\s*|```$", "", cleaned, flags=re.MULTILINE).strip()

        match = re.search(r"\{.*\}", cleaned, re.DOTALL)
        if not match:
            raise ValidationError("AI response did not contain valid JSON.")

        try:
            parsed = json.loads(match.group(0))
        except json.JSONDecodeError as exc:
            logger.error("Gemini JSON decode failed: %s | raw=%s", exc, raw_text[:500])
            raise ValidationError("Failed to parse AI response as JSON.")

        missing = [key for key in self.REQUIRED_KEYS if key not in parsed]
        if missing:
            raise ValidationError(f"AI response missing required fields: {missing}")

        for score_field in self.SCORE_FIELDS:
            parsed[score_field] = self._clamp_int(parsed.get(score_field, 0))

        for list_field in self.LIST_FIELDS:
            if not isinstance(parsed.get(list_field), list):
                parsed[list_field] = []

        for text_field in ("experience", "education"):
            if not isinstance(parsed.get(text_field), str):
                parsed[text_field] = str(parsed.get(text_field, "") or "")

        return parsed

    @staticmethod
    def _clamp_int(value: Any) -> int:
        try:
            value = int(round(float(value)))
        except (TypeError, ValueError):
            value = 0
        return max(0, min(100, value))

    def analyze(self, resume_text: str, job: Any) -> Dict[str, Any]:
        prompt = self._build_prompt(resume_text, job)
        try:
            response = self.model.generate_content(prompt)
            raw_text = response.text
        except Exception as exc:
            logger.exception("Gemini API call failed.")
            raise ValidationError(f"AI analysis failed: {exc}")

        parsed = self._parse_response(raw_text)
        parsed["_raw_response"] = raw_text
        return parsed

class SkillMatchingService:
    """Stage 1 — Role-based mandatory skill filtering."""

    @staticmethod
    def _normalize(skills: List[str]) -> List[str]:
        return [str(s).strip().lower() for s in skills if str(s).strip()]

    @staticmethod
    def check_mandatory_skills(job_skills: List[str], resume_skills: List[str]) -> Dict[str, Any]:
        required = SkillMatchingService._normalize(job_skills)
        found = SkillMatchingService._normalize(resume_skills)

        matched = [skill for skill in required if skill in found]
        missing = [skill for skill in required if skill not in found]

        passed = len(missing) == 0 if required else True

        return {
            "passed": passed,
            "matched_skills": matched,
            "missing_skills": missing,
            "match_percentage": round((len(matched) / len(required)) * 100, 2) if required else 100.0,
        }


class ResumeScoreService:
    """Stage 2 — Weighted overall score calculation.

    All sub-scores (skills_match, experience, education, projects,
    certifications, ats_formatting, keyword_match) are produced by Gemini.
    This service ONLY validates ranges and applies fixed weights to compute
    the final overall_score — it does not derive or estimate any sub-score.
    """

    WEIGHTS = {
        "skills_match_score": 0.30,
        "experience_score": 0.15,
        "education_score": 0.10,
        "projects_score": 0.10,
        "certifications_score": 0.05,
        "ats_score": 0.15,
        "keyword_match": 0.15,
    }

    REQUIRED_SCORE_FIELDS = list(WEIGHTS.keys())

    @classmethod
    def _validate_scores(cls, analysis_data: Dict[str, Any]) -> Dict[str, float]:
        validated: Dict[str, float] = {}
        for field in cls.REQUIRED_SCORE_FIELDS:
            value = analysis_data.get(field)
            if value is None:
                raise ValidationError(f"AI response missing required score field: {field}")
            try:
                value = float(value)
            except (TypeError, ValueError):
                raise ValidationError(f"AI score field '{field}' is not numeric.")
            if not (0 <= value <= 100):
                raise ValidationError(f"AI score field '{field}' out of range 0-100: {value}")
            validated[field] = value
        return validated

    @classmethod
    def calculate_score(
        cls,
        analysis_data: Dict[str, Any],
        resume_threshold: float,
    ) -> Dict[str, Any]:
        scores = cls._validate_scores(analysis_data)

        overall_score = sum(scores[field] * weight for field, weight in cls.WEIGHTS.items())
        overall_score = round(overall_score, 2)

        return {
            "skills_match_score": scores["skills_match_score"],
            "experience_score": scores["experience_score"],
            "education_score": scores["education_score"],
            "projects_score": scores["projects_score"],
            "certifications_score": scores["certifications_score"],
            "ats_formatting_score": scores["ats_score"],
            "keyword_match_score": scores["keyword_match"],
            "overall_score": overall_score,
            "threshold_used": float(resume_threshold),
            "is_shortlisted": overall_score >= float(resume_threshold),
        }

class ResumeAnalysisService:
    """Orchestrates parsing + AI analysis + persistence of ResumeAnalysis."""

    @staticmethod
    def parse_resume(resume: "Resume") -> str:
        file_path = resume.file.path
        text = ResumeParserService.extract_text(file_path, resume.file_type)
        resume.parsed_text = text
        resume.is_parsed = True
        resume.parse_error = None
        resume.save(update_fields=["parsed_text", "is_parsed", "parse_error"])
        return text

    @staticmethod
    def run_ai_analysis(resume_text: str, job: Any) -> Dict[str, Any]:
        analyzer = GeminiResumeAnalyzer()
        return analyzer.analyze(resume_text, job)

    @staticmethod
    def persist_analysis(
        application: "Application",
        analysis_data: Dict[str, Any],
        mandatory_check: Dict[str, Any],
    ) -> "ResumeAnalysis":
        from .models import ResumeAnalysis

        analysis, _ = ResumeAnalysis.objects.update_or_create(
            application=application,
            defaults={
                "skills": analysis_data.get("skills", []),
                "matched_skills": mandatory_check.get("matched_skills", []),
                "missing_skills": mandatory_check.get("missing_skills", []),
                "experience_summary": analysis_data.get("experience", ""),
                "education_summary": analysis_data.get("education", ""),
                "projects": analysis_data.get("projects", []),
                "certifications": analysis_data.get("certifications", []),
                "keyword_match": analysis_data.get("keyword_match", 0),
                "ats_score": analysis_data.get("ats_score", 0),
                "strengths": analysis_data.get("strengths", []),
                "weaknesses": analysis_data.get("weaknesses", []),
                "improvement_report": analysis_data.get("improvement_report", []),
                "raw_ai_response": {"text": analysis_data.get("_raw_response", "")},
                "mandatory_skills_passed": mandatory_check.get("passed", False),
            },
        )
        return analysis

    @staticmethod
    def persist_score(
        application: "Application",
        score_data: Dict[str, Any],
    ) -> "ResumeScore":
        from .models import ResumeScore

        score, _ = ResumeScore.objects.update_or_create(
            application=application,
            defaults=score_data,
        )
        return score


class ApplicationService:
    """Coordinates the full application -> screening -> shortlisting pipeline."""

    @staticmethod
    @transaction.atomic
    def apply_for_job(data: Dict[str, Any]) -> "Application":
        from apps.jobs.models import Job
        from .models import Application, Candidate, Resume

        job = Job.objects.filter(id=data["job_id"]).first()
        if not job:
            raise NotFound("Job not found.")
        if job.status != Job.Status.OPEN:
            raise ValidationError("This job is not currently accepting applications.")

        candidate, _ = Candidate.objects.update_or_create(
            email=data["email"],
            defaults={
                "first_name": data["first_name"],
                "last_name": data["last_name"],
                "phone": data["phone"],
                "location": data.get("location", ""),
                "linkedin_url": data.get("linkedin_url") or None,
                "portfolio_url": data.get("portfolio_url") or None,
            },
        )

        if Application.objects.filter(job=job, candidate=candidate).exists():
            raise ValidationError("You have already applied for this job.")

        resume_file = data["resume_file"]
        file_ext = resume_file.name.split(".")[-1].upper()
        resume = Resume.objects.create(
            candidate=candidate,
            file=resume_file,
            file_name=resume_file.name,
            file_type=file_ext,
        )

        application = Application.objects.create(
            job=job,
            candidate=candidate,
            resume=resume,
            status=Application.Status.APPLIED,
            cover_note=data.get("cover_note", ""),
        )

        ApplicationHistoryService.log(application, Application.Status.APPLIED, "Application submitted.")

        return application

    @staticmethod
    @transaction.atomic
    def process_application(application_id: int, user=None) -> "Application":
        """Runs the full screening pipeline: parse -> AI analyze -> mandatory check -> score -> decide."""
        from .models import Application

        application = ApplicationService.get_application(application_id, user)
        job = application.job
        resume = application.resume

        if resume is None:
            raise ValidationError("No resume attached to this application.")

        application.status = Application.Status.PROCESSING
        application.save(update_fields=["status", "updated_at"])
        ApplicationHistoryService.log(application, Application.Status.PROCESSING, "Screening started.")

        try:
            resume_text = ResumeAnalysisService.parse_resume(resume)
        except ValidationError as exc:
            resume.parse_error = str(exc)
            resume.is_parsed = False
            resume.save(update_fields=["parse_error", "is_parsed"])
            application.status = Application.Status.FAILED
            application.save(update_fields=["status", "updated_at"])
            ApplicationHistoryService.log(application, Application.Status.FAILED, f"Resume parsing failed: {exc}")
            raise

        try:
            analysis_data = ResumeAnalysisService.run_ai_analysis(resume_text, job)
        except ValidationError as exc:
            application.status = Application.Status.FAILED
            application.save(update_fields=["status", "updated_at"])
            ApplicationHistoryService.log(application, Application.Status.FAILED, f"AI analysis failed: {exc}")
            raise

        job_skills_required = job.skills_required if isinstance(job.skills_required, list) else []
        mandatory_check = SkillMatchingService.check_mandatory_skills(
            job_skills_required, analysis_data.get("skills", [])
        )

        ResumeAnalysisService.persist_analysis(application, analysis_data, mandatory_check)

        if not mandatory_check["passed"]:
            application.status = Application.Status.REJECTED_MANDATORY_SKILLS
            application.save(update_fields=["status", "updated_at"])
            ApplicationHistoryService.log(
                application,
                Application.Status.REJECTED_MANDATORY_SKILLS,
                f"Missing mandatory skills: {mandatory_check['missing_skills']}",
            )
            # Still persist a score record for reporting/UI consistency
            score_data = ResumeScoreService.calculate_score(
                analysis_data, float(job.resume_threshold)
            )
            score_data["is_shortlisted"] = False
            ResumeAnalysisService.persist_score(application, score_data)
            return application

        threshold = float(job.resume_threshold)
        score_data = ResumeScoreService.calculate_score(analysis_data, threshold)
        ResumeAnalysisService.persist_score(application, score_data)

        if score_data["is_shortlisted"]:
            application.status = Application.Status.SHORTLISTED
            remarks = f"Resume score {score_data['overall_score']} >= threshold {threshold}."
        else:
            application.status = Application.Status.REJECTED
            remarks = f"Resume score {score_data['overall_score']} < threshold {threshold}."

        application.save(update_fields=["status", "updated_at"])
        ApplicationHistoryService.log(application, application.status, remarks)

        return application

    @staticmethod
    def get_application(application_id: int, user=None) -> "Application":
        from .models import Application

        queryset = (
            Application.objects.select_related("job", "candidate", "resume", "analysis", "score")
            .prefetch_related("history")
        )
        # `user=None` is used for the internal candidate-submission flow
        # (ApplyForJobView), where there's no HR caller to scope against.
        # Every HR-facing call site passes `user` and gets tenant-scoped.
        if user is not None:
            queryset = queryset.filter(job__company=user.company)

        application = queryset.filter(id=application_id).first()
        if not application:
            raise NotFound("Application not found.")
        return application

    @staticmethod
    def list_applications(query_params: Dict[str, Any], user):
        from .models import Application

        queryset = (
            Application.objects.select_related("job", "candidate", "score")
            .filter(job__company=user.company)
        )

        job_id = query_params.get("job_id")
        if job_id:
            queryset = queryset.filter(job_id=job_id)

        status_filter = query_params.get("status")
        if status_filter:
            queryset = queryset.filter(status=status_filter)

        candidate_email = query_params.get("candidate_email")
        if candidate_email:
            queryset = queryset.filter(candidate__email__iexact=candidate_email)

        search = query_params.get("search")
        if search:
            queryset = queryset.filter(
                Q(candidate__first_name__icontains=search)
                | Q(candidate__last_name__icontains=search)
                | Q(candidate__email__icontains=search)
                | Q(job__title__icontains=search)
            )

        ordering = query_params.get("ordering")
        if ordering:
            queryset = queryset.order_by(ordering)

        return queryset

    @staticmethod
    @transaction.atomic
    def update_status(application_id: int, new_status: str, remarks: str = "", user=None) -> "Application":
        application = ApplicationService.get_application(application_id, user)
        application.status = new_status
        application.save(update_fields=["status", "updated_at"])
        ApplicationHistoryService.log(application, new_status, remarks or "Status updated by recruiter.")
        return application

    @staticmethod
    @transaction.atomic
    def update_recruiter_notes(application_id: int, notes: str, user=None) -> "Application":
        application = ApplicationService.get_application(application_id, user)
        application.recruiter_notes = notes
        application.save(update_fields=["recruiter_notes", "updated_at"])
        return application


class ApplicationHistoryService:
    @staticmethod
    def log(application: "Application", status: str, remarks: str = "") -> "ApplicationHistory":
        from .models import ApplicationHistory

        return ApplicationHistory.objects.create(
            application=application, status=status, remarks=remarks
        )

class CandidateDashboardService:
    """Candidate self-service views.

    NOTE: Identity resolution is centralized in `_resolve_candidate`.
    Today it resolves by verified email (current API contract, unchanged).
    When candidate authentication (JWT) is introduced, only this method
    needs to change — e.g. resolve via `request.user.candidate_profile`
    instead of an email lookup. All other methods in this service call
    `_resolve_candidate` and never query Candidate directly, so no other
    code will need to change.
    """

    @staticmethod
    def _resolve_candidate(email: str) -> "Candidate":
        """Single seam for candidate identity resolution.

        Current: resolves by verified email (query param).
        Future: swap this implementation to resolve from an authenticated
        JWT-bound candidate user without changing any calling method.
        """
        from .models import Candidate

        candidate = Candidate.objects.filter(email__iexact=email).first()
        if not candidate:
            raise NotFound("No candidate found with this email.")
        return candidate

    @staticmethod
    def get_candidate_by_email(email: str) -> "Candidate":
        # Kept for backward compatibility with any existing direct callers.
        return CandidateDashboardService._resolve_candidate(email)

    @staticmethod
    def get_dashboard(email: str) -> Dict[str, Any]:
        from .models import Application

        candidate = CandidateDashboardService._resolve_candidate(email)

        applications = (
            Application.objects.select_related("job", "job__company", "score")
            .filter(candidate=candidate)
            .order_by("-applied_at")
        )
        resumes = candidate.resumes.all().order_by("-uploaded_at")

        shortlisted = applications.filter(status=Application.Status.SHORTLISTED).count()
        rejected = applications.filter(
            status__in=[Application.Status.REJECTED, Application.Status.REJECTED_MANDATORY_SKILLS]
        ).count()
        in_progress = applications.filter(
            status__in=[Application.Status.APPLIED, Application.Status.PROCESSING, Application.Status.UNDER_REVIEW]
        ).count()

        return {
            "profile": candidate,
            "total_applications": applications.count(),
            "shortlisted_count": shortlisted,
            "rejected_count": rejected,
            "in_progress_count": in_progress,
            "resumes": resumes,
            "applications": applications,
        }

    @staticmethod
    def get_application_history(email: str):
        from .models import Application

        candidate = CandidateDashboardService._resolve_candidate(email)
        return (
            Application.objects.select_related("job", "job__company", "score")
            .filter(candidate=candidate)
            .order_by("-applied_at")
        )

    @staticmethod
    def get_application_status(email: str, application_id: int) -> "Application":
        from .models import Application

        candidate = CandidateDashboardService._resolve_candidate(email)
        application = Application.objects.filter(id=application_id, candidate=candidate).first()
        if not application:
            raise NotFound("Application not found for this candidate.")
        return application

    @staticmethod
    def get_resume_history(email: str):
        candidate = CandidateDashboardService._resolve_candidate(email)
        return candidate.resumes.all().order_by("-uploaded_at")

    @staticmethod
    def get_application_timeline(email: str, application_id: int):
        application = CandidateDashboardService.get_application_status(email, application_id)
        return application.history.all().order_by("-changed_at")

class CandidateDocumentService:
    @staticmethod
    @transaction.atomic
    def upload_document(data: Dict[str, Any]) -> "CandidateDocument":
        from .models import Application, CandidateDocument

        candidate = CandidateDashboardService._resolve_candidate(data["email"])

        application = None
        application_id = data.get("application_id")
        if application_id:
            application = Application.objects.filter(id=application_id, candidate=candidate).first()
            if not application:
                raise NotFound("Application not found for this candidate.")

        file_obj = data["file"]

        duplicate_exists = CandidateDocument.objects.filter(
            candidate=candidate,
            application=application,
            document_type=data["document_type"],
            file_name=file_obj.name,
        ).exists()
        if duplicate_exists:
            raise ValidationError(
                "An identical document (same type and file name) has already been uploaded."
            )

        document = CandidateDocument.objects.create(
            candidate=candidate,
            application=application,
            document_type=data["document_type"],
            file=file_obj,
            file_name=file_obj.name,
        )
        return document

    @staticmethod
    def list_documents(email: str):
        candidate = CandidateDashboardService._resolve_candidate(email)
        return candidate.documents.all().order_by("-uploaded_at")